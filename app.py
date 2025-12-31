import streamlit as st
from groq import Groq
from supabase import create_client, Client
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64

# --- 1. CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Maestro Estratega NEM", page_icon="🍎", layout="wide")

# --- 2. CONEXIÓN SILENCIOSA A BASE DE DATOS ---
def conectar_db():
    try:
        # Aquí el código intenta leer las llaves
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_KEY"]
        return create_client(url, key)
    except:
        # Si NO encuentra las llaves, devuelve "None" en lugar de un error rojo
        return None

supabase = conectar_db()

# --- INICIALIZACIÓN DE HISTORIAL DE SESIÓN ---
if 'historial_planeaciones' not in st.session_state:
    st.session_state.historial_planeaciones = []

# --- 3. FUNCIONES DE EXPORTACIÓN (SIN RECORTES) ---
def crear_word(contenido, nombre_proyecto):
    doc = Document()
    # Título principal centrado
    titulo = doc.add_heading(nombre_proyecto, 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Procesamiento de líneas para mantener formato
    for linea in contenido.split('\n'):
        if linea.strip():
            p = doc.add_paragraph(linea)
            p.style.font.size = Pt(11)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def generar_pdf_html(contenido, nombre_proyecto):
    # Formato de impresión automático que respeta tablas
    html = f"""
    <html>
    <head>
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;700&display=swap');
            body {{ font-family: 'Inter', sans-serif; padding: 40px; color: #1e293b; }}
            h1 {{ color: #1e3a8a; text-align: center; font-size: 28px; margin-bottom: 20px; border-bottom: 4px solid #1e3a8a; padding-bottom: 10px; }}
            table {{ border-collapse: collapse; width: 100%; margin-top: 20px; font-size: 12px; }}
            th, td {{ border: 1px solid #cbd5e1; padding: 12px; text-align: left; }}
            th {{ background-color: #f8fafc; color: #1e3a8a; font-weight: bold; }}
            .footer-pdf {{ margin-top: 30px; text-align: center; font-size: 10px; color: #94a3b8; }}
        </style>
    </head>
    <body>
        <h1>{nombre_proyecto}</h1>
        <div>{contenido.replace('|', '').replace('-', '')}</div>
        <div class='footer-pdf'>Generado por Maestro Estratega NEM - 2025</div>
    </body>
    </html>
    """
    b64 = base64.b64encode(html.encode()).decode()
    return f'<a href="data:text/html;base64,{b64}" download="{nombre_proyecto}.html" style="text-decoration:none;"><button style="width:100%; cursor:pointer; background: linear-gradient(45deg, #ef4444, #dc2626); color:white; padding:15px; border:none; border-radius:12px; font-weight:bold; font-size:16px; box-shadow: 0 4px 15px rgba(220, 38, 38, 0.3);">📄 DESCARGAR PDF AUTOMÁTICO (IMPRIMIR)</button></a>'

# --- 4. FUNCIONES DE LÓGICA DE USUARIO ---
def registrar_usuario(email, password, plan_elegido="Pendiente"):
    try:
        data = {
            "email": str(email), 
            "password": str(password), 
            "plan": plan_elegido, 
            "fecha_registro": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        supabase.table("usuarios").insert(data).execute()
        return True
    except: return False

def obtener_usuario(email):
    try:
        res = supabase.table("usuarios").select("*").eq("email", email).execute()
        return res.data[0] if res.data else None
    except: return None

# --- 5. ESTILOS CSS DE ALTO IMPACTO (DASHBOARD APP MODERNA) ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;800&display=swap');
    html, body, [class*="st-"] { font-family: 'Inter', sans-serif; }
    
    .hero-section {
        background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 100%);
        padding: 60px 30px; border-radius: 30px; color: white; text-align: center;
        margin-bottom: 40px; box-shadow: 0 20px 40px rgba(0,0,0,0.1);
    }
    
    .app-card {
        background: white; padding: 30px; border-radius: 24px; border: 1px solid #e2e8f0;
        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.05); margin-bottom: 25px;
    }
    
    .price-card-premium {
        background: white; padding: 40px; border-radius: 30px; border: 2px solid #3b82f6;
        text-align: center; transition: 0.4s;
    }
    .price-card-premium:hover { transform: translateY(-10px); box-shadow: 0 25px 50px -12px rgba(0, 0, 0, 0.15); }
    
    .stButton>button { border-radius: 12px; font-weight: 600; padding: 12px 24px; }
    .footer { text-align: center; padding: 50px; color: #64748b; font-size: 0.9rem; border-top: 1px solid #e2e8f0; }
    </style>
    """, unsafe_allow_html=True)

# --- 6. NAVEGACIÓN ---
if 'user' not in st.session_state:
    with st.sidebar:
        st.markdown("<h1 style='text-align:center;'>🧭 Menú</h1>", unsafe_allow_html=True)
        choice = st.radio("", ["🏠 Inicio", "📖 Acerca de", "📩 Contacto", "📝 Registrarse", "🔑 Iniciar Sesión"], label_visibility="collapsed")
        st.divider()
        st.image("https://via.placeholder.com/150x50?text=LOGO+NEM", use_container_width=True)

    if choice == "🏠 Inicio":
        st.markdown("<div class='hero-section'><h1>Maestro Estratega NEM</h1><p>Potenciando la educación con Inteligencia Artificial de Vanguardia</p></div>", unsafe_allow_html=True)
        st.markdown("### 📺 Video Promocional")
        st.video("https://www.youtube.com/watch?v=dQw4w9WgXcQ")
        st.divider()
        c1, c2, c3 = st.columns(3)
        with c1: st.markdown("<div class='app-card'><h3>🚀 Ahorro Real</h3><p>Genera propuestas pedagógicas en menos de 10 segundos y recupera tu tiempo personal.</p></div>", unsafe_allow_html=True)
        with c2: st.markdown("<div class='app-card'><h3>📚 Libros 2024</h3><p>Vinculación exacta con los Proyectos de Aula, Escolares y Comunitarios de la SEP.</p></div>", unsafe_allow_html=True)
        with c3: st.markdown("<div class='app-card'><h3>✅ Alineación NEM</h3><p>Basado fielmente en los Campos Formativos y Ejes Articuladores del Programa Sintético.</p></div>", unsafe_allow_html=True)

    elif choice == "📖 Acerca de":
        st.markdown("<h1 style='text-align: center;'>📖 Nuestra Historia y Misión</h1>", unsafe_allow_html=True)
        st.markdown("<div class='app-card'>", unsafe_allow_html=True)
        st.markdown("""
        **Maestro Estratega NEM** nació en el corazón de las aulas mexicanas. No somos solo una plataforma tecnológica; somos un equipo de docentes y desarrolladores que entendemos que el tiempo frente al grupo es lo más valioso, pero que la carga administrativa a menudo nos lo roba.
        
        Nuestra misión es democratizar el acceso a la **Inteligencia Artificial** de última generación para todos los docentes de México. Queremos que la implementación de la **Nueva Escuela Mexicana (NEM)** no sea una carga burocrática, sino una oportunidad para innovar y conectar mejor con nuestros alumnos.
        
        Estamos comprometidos con la actualización constante, siguiendo los lineamientos de la **Secretaría de Educación Pública (SEP)** y asegurando que cada planeación generada sea una herramienta real de transformación pedagógica.
        """)
        st.image("https://images.unsplash.com/photo-1524178232363-1fb2b075b655?auto=format&fit=crop&w=1000", caption="Innovación para el magisterio mexicano")
        st.markdown("</div>", unsafe_allow_html=True)

    elif choice == "📝 Registrarse":
        with st.form("reg_form"):
            st.subheader("Crea tu cuenta")
            email = st.text_input("Correo electrónico"); pw = st.text_input("Contraseña", type="password")
            if st.form_submit_button("Siguiente Paso"):
                st.session_state.temp_email, st.session_state.temp_pw, st.session_state.show_options = email, pw, True
        
        if st.session_state.get('show_options'):
            st.markdown("### 🎯 Selecciona tu Plan de Acceso")
            c1, c2, c3 = st.columns(3)
            with c1: 
                if st.button("🎁 Cortesía"): registrar_usuario(st.session_state.temp_email, st.session_state.temp_pw, "Cortesía"); st.rerun()
            with c2:
                if st.button("📅 Plan Anual"): registrar_usuario(st.session_state.temp_email, st.session_state.temp_pw, "Pendiente"); st.rerun()
            with c3:
                if st.button("💎 Plan 3 Años"): registrar_usuario(st.session_state.temp_email, st.session_state.temp_pw, "Pendiente"); st.rerun()

    elif choice == "🔑 Iniciar Sesión":
        with st.form("login_form"):
            e = st.text_input("Email"); p = st.text_input("Password", type="password")
            if st.form_submit_button("Entrar al Sistema"):
                u = obtener_usuario(e)
                if u and u['password'] == p: st.session_state.user = u; st.rerun()

else:
    # --- PANEL DE CONTROL (LOGUEADO) ---
    with st.sidebar:
        st.markdown(f"### Maestro:<br><span style='color:#3b82f6;'>{st.session_state.user['email']}</span>", unsafe_allow_html=True)
        st.write(f"Plan Actual: **{st.session_state.user['plan']}**")
        if st.button("Cerrar Sesión", use_container_width=True): del st.session_state.user; st.rerun()

    if st.session_state.user['plan'] == "Pendiente":
        st.title("💎 Activa tu Suscripción Profesional")
        cp1, cp2 = st.columns(2)
        with cp1:
            st.markdown("<div class='price-card-premium'><h3>PLAN ANUAL</h3><h1 style='font-size:3rem;'>$899</h1><p>Depósito OXXO/BBVA:<br><b>CLABE: 0123 4567 8901 2345 67</b></p></div>", unsafe_allow_html=True)
            st.link_button("🔥 Pagar con Stripe", "https://buy.stripe.com/TU_LINK", use_container_width=True)
        with cp2:
            st.markdown("<div class='price-card-premium'><h3>PLAN 3 AÑOS</h3><h1 style='font-size:3rem;'>$1,999</h1><p>Transferencia SPEI:<br><b>CLABE: 0123 4567 8901 2345 67</b></p></div>", unsafe_allow_html=True)
            st.link_button("🔵 Pagar con PayPal", "https://paypal.me/USUARIO", use_container_width=True)
    else:
        st.markdown("## 🤖 Estación de Planeación Inteligente NEM")
        
        # HISTORIAL DE SESIÓN (RESTAURADO)
        if st.session_state.historial_planeaciones:
            with st.expander("📂 Recuperar planeaciones anteriores (Sesión Actual)"):
                for idx, plan in enumerate(reversed(st.session_state.historial_planeaciones)):
                    if st.button(f"Recuperar: {plan['nombre']} ({plan['fecha']})", key=f"hist_btn_{idx}"):
                        st.session_state.resultado = plan['contenido']
                        st.session_state.nombre_p = plan['nombre']
                        st.rerun()

        # FORMULARIO COMPLETO (RESTAURADO)
        with st.form("planeacion_nem_form"):
            st.markdown("<div class='app-card'>", unsafe_allow_html=True)
            col1, col2, col3 = st.columns(3)
            with col1:
                fase = st.selectbox("Fase / Nivel", ["Fase 3", "Fase 4", "Fase 5", "Fase 6"])
                grado = st.selectbox("Grado Escolar", ["1°", "2°", "3°", "4°", "5°", "6°"])
                seccion = st.text_input("Sección", value="A", value="B")
            with col2:
                campo = st.selectbox("Campo Formativo", ["Lenguajes", "Saberes y P. Científico", "Ética, Nat. y Soc.", "De lo Humano y lo Com."])
                escenario = st.selectbox("Escenario", ["Aula", "Escolar", "Comunitario"])
            with col3:
                duracion = st.select_slider("Temporalidad", options=["1 día", "3 días", "1 semana", "2 semanas", "1 mes"])
                ejes = st.multiselect("Ejes Articuladores", ["Inclusión", "Pensamiento Crítico", "Interculturalidad Crítica", 
                                                             "Igualdad de Género", "Vida Saludable", 
                                                             "Apropiación de las Culturas a través de la Lectura y la Escritura", "Artes y Experiencias Estéticas"])
            
            tema = st.text_area("Nombre del Proyecto o Problemática Central:")
            
            if st.form_submit_button("✨ GENERAR PLANEACIÓN INTEGRAL"):
                if tema:
                    with st.spinner("Buscando en Programa Sintético y vinculando Libros SEP 2024..."):
                        client = Groq(api_key=st.secrets["GROQ_API_KEY"])
                        prompt = f"""
                        Eres experto en la NEM 2024 (SEP México). Genera planeación para {grado} de {fase}, Sección {seccion}.
                        Campo: {campo}, Escenario: {escenario}. DURACIÓN: {duracion}.
                        Tema: {tema}.
                        
                        REQUISITOS ESTRICTOS:
                        1. TODO EL RESULTADO EN TABLAS DE MARKDOWN.
                        2. TABLA 1: Datos Generales, Contenido oficial y PDA (Buscados en el Programa Sintético).
                        3. TABLA 2: Vinculación con LIBROS SEP 2024 vigentes: Nombre del Proyecto y PÁGINAS EXACTAS de acuerdo al grado {grado}.
                        4. TABLA 3: Secuencia Didáctica por sesión (45-50 min):
                           - INICIO (10 min): Actividad y Materiales.
                           - DESARROLLO (30 min): Actividad central (Metodologías NEM) y Materiales específicos.
                           - CIERRE (10 min): Metacognición y Evaluación.
                        Asegúrate de que el número de sesiones coincida con {duracion}.
                        """
                        res = client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role": "user", "content": prompt}])
                        st.session_state.resultado = res.choices[0].message.content
                        st.session_state.nombre_p = f"Proyecto: {tema[:40]}"
                        st.session_state.historial_planeaciones.append({"nombre": st.session_state.nombre_p, "contenido": st.session_state.resultado, "fecha": datetime.now().strftime("%H:%M")})
                        st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)

        if 'resultado' in st.session_state:
            st.markdown(f"<div class='app-card'><h2>📍 {st.session_state.nombre_p}</h2>", unsafe_allow_html=True)
            st.markdown(st.session_state.resultado)
            st.divider()
            d1, d2 = st.columns(2)
            with d1:
                st.download_button("📄 Descargar en Word (.docx)", crear_word(st.session_state.resultado, st.session_state.nombre_p), f"{st.session_state.nombre_p}.docx", use_container_width=True)
            with d2:
                st.markdown(generar_pdf_html(st.session_state.resultado, st.session_state.nombre_p), unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

st.markdown("<div class='footer'>© 2025 Maestro Estratega NEM | Matamoros, Tamaulipas.</div>", unsafe_allow_html=True)






